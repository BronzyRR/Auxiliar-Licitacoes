from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from ModuloGeral import *

def CM2IN(medida):
    return medida / 2.54


# Classe com todos os estilos necessários para a criação dos Anexos nos moldes dos exigidos pela prefeitura
class DocumentMultivendas:

    @staticmethod
    def criarComEstilo():
        documentoBase = Document()
        styles = documentoBase.styles

        # estilo para o título dos anexos
        estilo_titulo = styles.add_style("TitleAnexos", WD_STYLE_TYPE.PARAGRAPH)
        estilo_titulo.paragraph_format.space_after = Inches(CM2IN(0.50))
        estilo_titulo.paragraph_format.line_spacing = Inches(CM2IN(0.75))
        estilo_titulo.font.name = 'Arial'
        estilo_titulo.font.bold = True
        estilo_titulo.font.size = Pt(14)

        # estilo para o subtítulo dos anexos
        estilo_sub_titulo = styles.add_style("SubTitleAnexos", WD_STYLE_TYPE.PARAGRAPH)
        estilo_sub_titulo.paragraph_format.line_spacing = Inches(CM2IN(0.75))
        estilo_sub_titulo.font.name = 'Arial'
        estilo_sub_titulo.font.bold = False
        estilo_sub_titulo.font.size = Pt(12)

        # estilo para o corpo do texto
        estilo_corpo_texto = styles.add_style("TextoAnexos", WD_STYLE_TYPE.PARAGRAPH)
        estilo_corpo_texto.paragraph_format.first_line_indent = Inches(CM2IN(1.50))
        estilo_corpo_texto.paragraph_format.space_before = Inches(CM2IN(1.0))
        estilo_corpo_texto.paragraph_format.space_after = Inches(CM2IN(0.75))
        estilo_corpo_texto.paragraph_format.line_spacing = Inches(CM2IN(0.60))
        estilo_corpo_texto.font.name = 'Arial'
        estilo_corpo_texto.font.size = Pt(12)

        # estilo para a linha de assinatura
        estilo_assinatura = styles.add_style("TextoSemIdentacao", WD_STYLE_TYPE.PARAGRAPH)
        estilo_assinatura.paragraph_format.space_after = Inches(CM2IN(0.25))
        estilo_assinatura.paragraph_format.line_spacing = Inches(CM2IN(0.50))
        estilo_assinatura.font.name = 'Arial'
        estilo_assinatura.font.size = Pt(12)

        # estilo para a data
        estilo_data = styles.add_style("DataAnexos", WD_STYLE_TYPE.PARAGRAPH)
        estilo_data.paragraph_format.space_after = Inches(CM2IN(1.25))
        estilo_data.font.name = 'Arial'
        estilo_data.font.bold = False
        estilo_data.font.size = Pt(12)

        # estilo para o CD
        estilo_cd = styles.add_style("CD", WD_STYLE_TYPE.PARAGRAPH)
        estilo_cd.paragraph_format.line_spacing = 1
        estilo_cd.font.name = 'Arial'
        estilo_cd.font.bold = False
        estilo_cd.font.size = Pt(11)

        # estilo para o rodape
        estilo_data = styles.add_style("Rodape", WD_STYLE_TYPE.PARAGRAPH)
        estilo_data.font.name = 'Arial'
        estilo_data.font.size = Pt(9)

        # estilo para as alíneas
        estilo_alinea = styles.add_style("Alinea", WD_STYLE_TYPE.PARAGRAPH)
        estilo_alinea.font.name = 'Arial'
        estilo_alinea.font.size = Pt(10)
        estilo_alinea.paragraph_format.left_indent = Pt(50)

        # todos os documentos do anexo
        documentoBase.add_picture('configs/logo.png', width=Inches(CM2IN(10)))
        paragrafo_imagem = documentoBase.paragraphs[-1]
        paragrafo_imagem.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # criar o rodapé
        footer = documentoBase.sections[0].footer
        paragrafo_rodape = footer.paragraphs[0]
        rodape_run = paragrafo_rodape.add_run("")
        rodape_run.add_picture("configs/rodape.png", width=Inches(CM2IN(10)))
        paragrafo_rodape.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return documentoBase

    @staticmethod
    def criarComEstiloSemRodape():
        documentoBase = Document()
        styles = documentoBase.styles

        # estilo para o título dos anexos
        estilo_titulo = styles.add_style("TitleAnexos", WD_STYLE_TYPE.PARAGRAPH)
        estilo_titulo.paragraph_format.space_after = Inches(CM2IN(0.50))
        estilo_titulo.paragraph_format.line_spacing = Inches(CM2IN(0.75))
        estilo_titulo.font.name = 'Arial'
        estilo_titulo.font.bold = True
        estilo_titulo.font.size = Pt(14)

        # estilo para o subtítulo dos anexos
        estilo_sub_titulo = styles.add_style("SubTitleAnexos", WD_STYLE_TYPE.PARAGRAPH)
        estilo_sub_titulo.paragraph_format.line_spacing = Inches(CM2IN(0.75))
        estilo_sub_titulo.font.name = 'Arial'
        estilo_sub_titulo.font.bold = False
        estilo_sub_titulo.font.size = Pt(12)

        # estilo para o corpo do texto
        estilo_corpo_texto = styles.add_style("TextoAnexos", WD_STYLE_TYPE.PARAGRAPH)
        estilo_corpo_texto.paragraph_format.first_line_indent = Inches(CM2IN(1.50))
        estilo_corpo_texto.paragraph_format.space_before = Inches(CM2IN(1.0))
        estilo_corpo_texto.paragraph_format.space_after = Inches(CM2IN(0.75))
        estilo_corpo_texto.paragraph_format.line_spacing = Inches(CM2IN(0.60))
        estilo_corpo_texto.font.name = 'Arial'
        estilo_corpo_texto.font.size = Pt(12)

        # estilo para a linha de assinatura
        estilo_assinatura = styles.add_style("TextoSemIdentacao", WD_STYLE_TYPE.PARAGRAPH)
        estilo_assinatura.paragraph_format.space_after = Inches(CM2IN(0.25))
        estilo_assinatura.paragraph_format.line_spacing = Inches(CM2IN(0.50))
        estilo_assinatura.font.name = 'Arial'
        estilo_assinatura.font.size = Pt(12)

        # estilo para a data
        estilo_data = styles.add_style("DataAnexos", WD_STYLE_TYPE.PARAGRAPH)
        estilo_data.paragraph_format.space_after = Inches(CM2IN(1.25))
        estilo_data.font.name = 'Arial'
        estilo_data.font.bold = False
        estilo_data.font.size = Pt(12)

        # estilo para o CD
        estilo_cd = styles.add_style("CD", WD_STYLE_TYPE.PARAGRAPH)
        estilo_cd.paragraph_format.line_spacing = 1
        estilo_cd.font.name = 'Arial'
        estilo_cd.font.bold = False
        estilo_cd.font.size = Pt(11)

        # estilo para as alíneas
        estilo_alinea = styles.add_style("Alinea", WD_STYLE_TYPE.PARAGRAPH)
        estilo_alinea.font.name = 'Arial'
        estilo_alinea.font.size = Pt(10)
        estilo_alinea.paragraph_format.left_indent = Pt(50)

        # todos os documentos do anexo
        documentoBase.add_picture('configs/logo.png', width=Inches(CM2IN(10)))
        paragrafo_imagem = documentoBase.paragraphs[-1]
        paragrafo_imagem.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        return documentoBase


def criarAnexoV(razaoSocial, CNPJ, enderecoSede, data, nomeRepresentanteLegal, identidadeRepresentanteLegal,
                orgaoEmissor):
    anexo_v = DocumentMultivendas.criarComEstilo()

    # TÍTULO DO ANEXO V
    titulo = anexo_v.add_paragraph("DECLARAÇÃO DE INEXISTÊNCIA DE FATO IMPEDITIVO", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO PRINCIPAL DO ANEXO V
    texto_principal = anexo_v.add_paragraph("Na qualidade de representante legal da empresa ", style="TextoAnexos")
    texto_principal.add_run("{}".format(razaoSocial)).bold = True
    texto_principal.add_run(", inscrita no CNPJ sob n° ")
    texto_principal.add_run("{}".format(CNPJ)).bold = True
    texto_principal.add_run(", com sede localizada na ")
    texto_principal.add_run("{}".format(enderecoSede)).bold = True
    texto_principal.add_run(
        ", declaro, sob as penas da Lei que até a presente data, inexistem fatos impeditivos para sua habilitação"
        " no presente Processo Licitatório, ciente da obrigatoriedade de declarar as ocorrências posteriores.")
    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = anexo_v.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = anexo_v.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = anexo_v.add_paragraph("{}".format(nomeRepresentanteLegal), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = anexo_v.add_paragraph("RG: {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor),
                                     style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    anexo_v.save("(05) anexo_v.docx")


def criar_anexoVI(razaoSocial, CNPJ, enderecoSede, concorrencia, data, nomeRepresentanteLegal,
                  identidadeRepresentanteLegal, orgaoEmissor):
    anexo_vi = DocumentMultivendas.criarComEstilo()

    # TÍTULO DO ANEXO VI
    titulo = anexo_vi.add_paragraph("DECLARAÇÃO DE NÃO EMPREGO DE MENOR", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO PRINCIPAL DO ANEXO VI
    texto_principal = anexo_vi.add_paragraph("Na qualidade de representante legal da empresa ", style="TextoAnexos")
    texto_principal.add_run("{}".format(razaoSocial)).bold = True
    texto_principal.add_run(", inscrita no CNPJ sob o n°")
    texto_principal.add_run(" {}".format(CNPJ)).bold = True
    texto_principal.add_run(", com sede localizada na")
    texto_principal.add_run(" {}".format(enderecoSede)).bold = True
    texto_principal.add_run(", declaro, para efeito de habilitação na")
    texto_principal.add_run(" {}".format(concorrencia)).bold = True
    texto_principal.add_run(
        ", que não há, no quadro de pessoal da referida empresa, empregado(s) com menos de 18 (dezoito)"
        " anos em trabalho noturno, perigoso ou insalubre e de 16 (dezesseis) anos, em qualquer trabalho, salvo"
        " na condição de aprendiz, a partir de 14 (quatorze) anos, nos termos do inciso XXXIII do art 7° da Constituição Federal.")
    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = anexo_vi.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = anexo_vi.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = anexo_vi.add_paragraph("{}".format(nomeRepresentanteLegal), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = anexo_vi.add_paragraph("RG: {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor),
                                      style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    anexo_vi.save("(06) anexo_vi.docx")


def criar_anexoVII(razaoSocial, CNPJ, enderecoSede, concorrencia, data, nomeRepresentanteLegal,
                   identidadeRepresentanteLegal, orgaoEmissor):
    anexo_vii = DocumentMultivendas.criarComEstilo()

    # TÍTULO DO ANEXO VII
    titulo = anexo_vii.add_paragraph(
        "DECLARAÇÃO DE TRABALHO DE EMPREGADO SEM VÍNCULO COM A PREFEITURA MUNICIPAL DE BOA VISTA E RESPONSÁVEL PELA LICITAÇÃO",
        style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO PRINCIPAL DO ANEXO VII
    texto_principal = anexo_vii.add_paragraph("Na qualidade de representante legal da empresa", style="TextoAnexos")
    texto_principal.add_run(" {}".format(razaoSocial)).bold = True
    texto_principal.add_run(", inscrita no CNPJ sob o n°")
    texto_principal.add_run(" {}".format(CNPJ)).bold = True
    texto_principal.add_run(", com sede localizada na ")
    texto_principal.add_run(" {}".format(enderecoSede)).bold = True
    texto_principal.add_run(", declaro, para efeito de habilitação na")
    texto_principal.add_run(" {}".format(concorrencia)).bold = True
    texto_principal.add_run(
        ", que não há, no quadro de pessoal da referida empresa, empregado(s) que seja(m) cônjuge, companheiro"
        " ou parentes em linha reta, colateral ou por afinidade, até o terceiro grau, inclusive, de ocupantes"
        " de cargos de direção e assessoramento, de funcionários vinculados ao MUNICÍPIO DE BOA VISTA, nos termos"
        " do Art. 9°, inciso III, da Lei n° 8.666/93.")
    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = anexo_vii.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = anexo_vii.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = anexo_vii.add_paragraph("{}".format(nomeRepresentanteLegal), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = anexo_vii.add_paragraph("RG: {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor),
                                       style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    anexo_vii.save("(07) anexo_vii.docx")


def criar_anexoVIII(razaoSocial, CNPJ, data, nomeRepresentanteLegal, identidadeRepresentanteLegal, orgaoEmissor, CPF):
    anexo_viii = DocumentMultivendas.criarComEstilo()

    # TÍTULO DO ANEXO VIII
    titulo = anexo_viii.add_paragraph(
        "DECLARAÇÃO DE QUE A EMPRESA LICITANTE REALIZARÁ OS SERVIÇOS DESTE EDITAL DE ACORDO COM A LEGISLAÇÃO AMBIENTAL VIGENTE",
        style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO PRINCIPAL DO ANEXO VIII
    texto_principal = anexo_viii.add_paragraph("A empresa", style="TextoAnexos")
    texto_principal.add_run(" {}".format(razaoSocial)).bold = True
    texto_principal.add_run(", inscrita no CNPJ n°")
    texto_principal.add_run(" {}".format(CNPJ)).bold = True
    texto_principal.add_run(" por intermédio de seu representante legal o Sr ")
    texto_principal.add_run(" {}".format(nomeRepresentanteLegal)).bold = True
    texto_principal.add_run(", portador da Carteira de Identidade n°")
    texto_principal.add_run(" {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor)).bold = True
    texto_principal.add_run(" e do CPF n°")
    texto_principal.add_run(" {}".format(CPF)).bold = True
    texto_principal.add_run(
        ", DECLARA, que realizará os serviços objetos do presente edital em conformidade com a Legislação Ambiental Vigente.")
    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = anexo_viii.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = anexo_viii.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = anexo_viii.add_paragraph("{}".format(nomeRepresentanteLegal), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = anexo_viii.add_paragraph("RG: {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor),
                                        style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    anexo_viii.save("(08) anexo_viii.docx")


def criar_documento01_aceitacao(concorrencia, processo, nomeResponsavelTecnico, cpf_responsavel_tecnico, registroCREA,
                                data, enderecoSalvar):
    documento_01 = DocumentMultivendas.criarComEstilo()

    # TITULO DO DOCUMENTO 01 (DECLARAÇÃO DE AUTORIZAÇÃO DE INCLUSÃO NA EQUIPE TÉCNICA)
    titulo = documento_01.add_paragraph("DECLARAÇÃO DE AUTORIZAÇÃO DE INCLUSÃO NA EQUIPE TÉCNICA", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitulo_concorrencia = documento_01.add_paragraph("{}".format(concorrencia), style="SubTitleAnexos")
    subtitulo_concorrencia.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitulo_processo = documento_01.add_paragraph("{}".format(processo), style="SubTitleAnexos")
    subtitulo_processo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO PRINCIPAL DO DOCUMENTO 01
    texto_principal = documento_01.add_paragraph("Em atendimento ao previsto no edital da", style="TextoAnexos")
    texto_principal.add_run(" {}".format(concorrencia)).bold = True
    texto_principal.add_run(", Eu").bold = False
    texto_principal.add_run(" {}".format(nomeResponsavelTecnico)).bold = True
    texto_principal.add_run(" - Engenheiro Civil,")
    texto_principal.add_run(" CPF: {}".format(cpf_responsavel_tecnico)).bold = True
    texto_principal.add_run(" e portador da carteira de registro no CREA-RR número")
    texto_principal.add_run(" {}".format(registroCREA)).bold = True
    texto_principal.add_run(
        ", autorizo minha inclusão na equipe técnica que irá participar na execução da obra/serviços"
        " objeto da licitação em referência.")
    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento_01.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento_01.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO RESPONSÁVEL TÉCNICO
    linha_representante_legal = documento_01.add_paragraph("{}".format(nomeResponsavelTecnico),
                                                           style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # REGISTRO DO CREA DO RESPONSÁVEL TÉCNICO
    linha_rg = documento_01.add_paragraph("CREA n° {}".format(registroCREA),
                                          style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    documento_01.save("{}/(01) ACEITAÇÃO RESPONSÁVEL TÉCNICO.docx".format(enderecoSalvar))


def criar_documento02_equipe_tecnica(concorrencia, processo, nomeResponsavelTecnico, cpf_responsavel_tecnico,
                                     registroCREA, data,
                                     nomeRepresentanteLegal, cargo, enderecoSalvar):
    documento = DocumentMultivendas.criarComEstilo()

    # TITULO DO DOCUMENTO 02 (DECLARAÇÃO DE INDICAÇÃO DOS RESPONSÁVEIS TÉCNICOS)
    titulo = documento.add_paragraph("DECLARAÇÃO DE INDICAÇÃO DOS RESPONSÁVEIS TÉCNICOS", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitulo_concorrencia = documento.add_paragraph("{}".format(concorrencia), style="SubTitleAnexos")
    subtitulo_concorrencia.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitulo_processo = documento.add_paragraph("{}".format(processo), style="SubTitleAnexos")
    subtitulo_processo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO PRINCIPAL DO DOCUMENTO 02
    texto_principal = documento.add_paragraph("Declaramos, em atendimento ao previsto no Edital da",
                                              style="TextoAnexos")
    texto_principal.add_run(" {}".format(concorrencia)).bold = True
    texto_principal.add_run(", que o engenheiro civil abaixo pertencente ao quadro técnico")
    texto_principal.add_run(" é o nosso indicado").bold = True
    texto_principal.add_run(" como responsável pelos trabalhos objeto da licitação em referência.")
    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO RESPONSÁVEL TÉCNICO
    linha_responsavel_tecnico = documento.add_paragraph("{}".format(nomeResponsavelTecnico),
                                                        style="TextoSemIdentacao")
    linha_responsavel_tecnico.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_responsavel_tecnico.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # REGISTRO DO CREA DO RESPONSÁVEL TÉCNICO
    linha_crea = documento.add_paragraph("CREA n° {}".format(registroCREA),
                                         style="TextoSemIdentacao")
    linha_crea.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_crea.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    # REGISTRO DO CPF DO RESPONSÁVEL TÉCNICO
    linha_cpf_responsavel_tecnico = documento.add_paragraph("CPF: {}".format(cpf_responsavel_tecnico),
                                                            style="TextoSemIdentacao")
    linha_cpf_responsavel_tecnico.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_cpf_responsavel_tecnico.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    # LINHA PARA ASSINATURA DO REPRESENTANTE LEGAL
    linha_assinatura_02 = documento.add_paragraph("\n___________________________________", style="TextoSemIdentacao")
    linha_assinatura_02.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_nome_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA", style="TextoSemIdentacao")
    linha_nome_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_nome_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE LEGAL
    linha_representante_legal = documento.add_paragraph("{}".format(nomeRepresentanteLegal),
                                                        style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # CARGO DO REPRESENTANTE LEGAL
    linha_cargo_representante_legal = documento.add_paragraph("{}".format(cargo),
                                                              style="TextoSemIdentacao")
    linha_cargo_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_cargo_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    documento.save("{}/(02) DECLARAÇÃO INDICAÇÃO EQUIPE TÉCNICA.docx".format(enderecoSalvar))


def criar_documento03_pertence_quadro_empresa(concorrencia, processo, nomeResponsavelTecnico, cpf_responsavel_tecnico,
                                              registroCREA, data, nomeRepresentanteLegal, cargo, enderecoSalvar):
    documento = DocumentMultivendas.criarComEstilo()

    # TITULO DO DOCUMENTO 03 (DECLARAÇÃO QUE PERTENCE AO QUADRO DA EMPRESA)
    titulo = documento.add_paragraph("DECLARAÇÃO QUE PERTENCE AO QUADRO DA EMPRESA", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitulo_concorrencia = documento.add_paragraph("{}".format(concorrencia), style="SubTitleAnexos")
    subtitulo_concorrencia.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitulo_processo = documento.add_paragraph("{}".format(processo), style="SubTitleAnexos")
    subtitulo_processo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO PRINCIPAL DO DOCUMENTO 03
    texto_principal = documento.add_paragraph("Declaramos, em atendimento ao previsto no Edital da",
                                              style="TextoAnexos")
    texto_principal.add_run(" {}".format(concorrencia)).bold = True
    texto_principal.add_run(", que o engenheiro civil indicado abaixo pertence ao quadro técnico da empresa.")
    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO RESPONSÁVEL TÉCNICO
    linha_responsavel_tecnico = documento.add_paragraph("{}".format(nomeResponsavelTecnico),
                                                        style="TextoSemIdentacao")
    linha_responsavel_tecnico.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_responsavel_tecnico.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # REGISTRO DO CREA DO RESPONSÁVEL TÉCNICO
    linha_crea = documento.add_paragraph("CREA n° {}".format(registroCREA),
                                         style="TextoSemIdentacao")
    linha_crea.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_crea.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    # REGISTRO DO CPF DO RESPONSÁVEL TÉCNICO
    linha_cpf_responsavel_tecnico = documento.add_paragraph("CPF: {}".format(cpf_responsavel_tecnico),
                                                            style="TextoSemIdentacao")
    linha_cpf_responsavel_tecnico.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_cpf_responsavel_tecnico.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    # LINHA PARA ASSINATURA DO REPRESENTANTE LEGAL
    linha_assinatura_02 = documento.add_paragraph("\n___________________________________", style="TextoSemIdentacao")
    linha_assinatura_02.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_nome_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA", style="TextoSemIdentacao")
    linha_nome_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_nome_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE LEGAL
    linha_representante_legal = documento.add_paragraph("{}".format(nomeRepresentanteLegal),
                                                        style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # CARGO DO REPRESENTANTE LEGAL
    linha_cargo_representante_legal = documento.add_paragraph("{}".format(cargo),
                                                              style="TextoSemIdentacao")
    linha_cargo_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_cargo_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    documento.save(
        "{}/(03) DECLARAÇÃO PERTENCE AO QUADRO DA EMPRESA DO RESPONSAVEL TÉCNICO.docx".format(enderecoSalvar))


def criar_documento04_alvara(concorrencia, processo, nomeRepresentanteLegal, cargoRepresentanteLegal, objeto,
                             cnpj_empresa, enderecoSede,
                             identidadeRepresentanteLegal, orgaoEmissor, data, enderecoSalvar):
    documento = DocumentMultivendas.criarComEstilo()

    # TITULO DO DOCUMENTO 04 (DECLARAÇÃO DE DISPONIBILIDADE DE ALVARÁ DE FUNCIONAMENTO)
    titulo = documento.add_paragraph("DECLARAÇÃO DE DISPONIBILIDADE DE ALVARÁ DE FUNCIONAMENTO", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitulo_concorrencia = documento.add_paragraph("{}".format(concorrencia), style="SubTitleAnexos")
    subtitulo_concorrencia.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitulo_processo = documento.add_paragraph("{}".format(processo), style="SubTitleAnexos")
    subtitulo_processo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO PRINCIPAL DO DOCUMENTO 04
    texto_principal = documento.add_paragraph("A empresa ", style="TextoAnexos")
    texto_principal.add_run(" MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA.").bold = True
    texto_principal.add_run(", CNPJ {}, sediada na {}, Boa Vista - RR".format(cnpj_empresa, enderecoSede))
    texto_principal.add_run(" por intermédio de seu representante legal infra-assinado e para os fins do")
    texto_principal.add_run(" EDITAL DA {},".format(concorrencia)).bold = True
    texto_principal.add_run(" {}".format(processo))
    texto_principal.add_run(", OBJETO: {},".format(objeto))
    texto_principal.add_run(" DECLARA EXPRESSAMENTE QUE DISPÕE DE ALVARÁ DE FUNCIONAMENTO VÁLIDO.").bold = True

    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA DO REPRESENTANTE LEGAL
    linha_assinatura_02 = documento.add_paragraph("\n___________________________________", style="TextoSemIdentacao")
    linha_assinatura_02.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_nome_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA", style="TextoSemIdentacao")
    linha_nome_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_nome_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE LEGAL
    linha_representante_legal = documento.add_paragraph("{}".format(nomeRepresentanteLegal),
                                                        style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # CARGO DO REPRESENTANTE LEGAL
    linha_cargo_representante_legal = documento.add_paragraph("RG: {} {}".format(identidadeRepresentanteLegal,
                                                                                 orgaoEmissor),
                                                              style="TextoSemIdentacao")

    linha_cargo_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_cargo_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    documento.save("{}/(04) DECLARAÇÃO ALVARÁ DE FUNCIONAMENTO VÁLIDO.docx".format(enderecoSalvar))


def criar_documento05_total_conhecimento(razaoSocial, CNPJ, concorrencia, objeto, data, nomeRepresentanteLegal,
                                         identidadeRepresentanteLegal,
                                         orgaoEmissor, enderecoSalvar):
    documento = DocumentMultivendas.criarComEstilo()

    # TITULO DO DOCUMENTO
    titulo = documento.add_paragraph(
        "DECLARAÇÃO DE TOTAL CONHECIMENTO, ACEITAÇÃO E QUE DISPÕE DE EQUIPAMENTOS NECESSÁRIOS "
        "PARA A EXECUÇÃO DOS SERVIÇOS", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO PRINCIPAL DO ANEXO IV
    texto_principal = documento.add_paragraph("A empresa ", style="TextoAnexos")
    texto_principal.add_run('{}'.format(razaoSocial)).bold = True
    texto_principal.add_run(', CNPJ n° ')
    texto_principal.add_run('{}'.format(CNPJ)).bold = True
    texto_principal.add_run(', declara em atendimento ao Edital de ')
    texto_principal.add_run('{}'.format(concorrencia)).bold = True
    texto_principal.add_run(', que tem por objeto ')
    texto_principal.add_run('{}'.format(objeto)).bold = True
    texto_principal.add_run(
        ', que tem total conhecimento e aceitação das condições estipuladas no Edital e dispõe de equipamentos necessários'
        ' a execução dos serviços.')
    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = documento.add_paragraph("{}".format(nomeRepresentanteLegal), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = documento.add_paragraph("RG: {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor),
                                       style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    documento.save(
        "{}/(05) DECLARAÇÃO DE TOTAL CONHECIMENTO, ACEITAÇÃO E QUE DISPÕE DE EQUIPAMENTOS NECESSÁRIOS.docx".format(
            enderecoSalvar))


def criar_documento06_declaracao_diversas(razaoSocial, CNPJ, enderecoSede, concorrencia, processo, data,
                                          nomeRepresentanteLegal, identidadeRepresentanteLegal, orgaoEmissor,
                                          enderecoSalvar):
    documento = DocumentMultivendas.criarComEstilo()

    # TITULO DO DOCUMENTO
    titulo = documento.add_paragraph("DECLARAÇÕES DIVERSAS", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO DECLARAÇÃO CABEÇALHO
    texto_declaracao_cabecalho = documento.add_paragraph("A empresa", style="TextoAnexos")
    texto_declaracao_cabecalho.add_run(" {}".format(razaoSocial)).bold = True
    texto_declaracao_cabecalho.add_run(", CNPJ {}, sediada na {}, Boa Vista - RR,".format(CNPJ, enderecoSede))
    texto_declaracao_cabecalho.add_run(" por intermédio de seu representante legal, infra-assinado e para os fins do ")
    texto_declaracao_cabecalho.add_run(" EDITAL DA {} e {}".format(concorrencia, processo)).bold = True
    texto_declaracao_cabecalho.add_run(" realiza as seguintes declarações:")
    texto_declaracao_cabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DECLARAÇÃO 01
    texto_declaracao_01 = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao_01.add_run("DECLARAMOS").bold = True
    texto_declaracao_01.add_run(" que nossa Empresa não se encontra inadimplente, nem é objeto de qualquer restrição ou"
                                " nota desabonadora junto ao cadastro de fornecedores dessa Secretaria ou qualquer esfera"
                                " da administração pública direta (Federal, Estadual ou Municipal) e indireta.")
    texto_declaracao_01.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DECLARAÇÃO 02
    texto_declaracao_02 = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao_02.add_run("DECLARAMOS").bold = True
    texto_declaracao_02.add_run(", sob as penas da lei, que inexiste até a presente data, fatos impeditivos para "
                                "sua habilitação e contratação no processo licitatório em referência, conforme artigo"
                                " 32 - da lei 8.666/93 e IM-MARE 05/95, ciente da obrigatoriedade de declarar "
                                "ocorrências posteriores.")
    texto_declaracao_02.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DECLARAÇÃO 03
    texto_declaracao_03 = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao_03.add_run("DECLARAMOS").bold = True
    texto_declaracao_03.add_run(" que temos conhecimento e aceitamos de forma integral e irretratável os termos deste"
                                " edital, seus anexos e instruções, bem como a observação dos regulamentos administrativos"
                                " e das normas técnicas gerais ou especiais aplicáveis.")
    texto_declaracao_03.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DECLARAÇÃO 04
    texto_declaracao_04 = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao_04.add_run("DECLARAMOS").bold = True
    texto_declaracao_04.add_run(
        " que nossa empresa não possui em seu quadro funcional empregados com idade inferior a 18 (dezoito)"
        " anos em trabalho noturno, perigoso ou insalubre e de 16 (dezesseis) anos em qualquer trabalho, "
        "salvo na condição de aprendiz, a partir de 14 (quatorze) anos (Lei 9.854/99 e Art 7°, XXXIII da CF).")
    texto_declaracao_04.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DECLARAÇÃO 05
    texto_declaracao_05 = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao_05.add_run("DECLARAMOS").bold = True
    texto_declaracao_05.add_run(
        " que os sócios da proponente, administradores, empregados e controladores não são servidores"
        " públicos, dirigentes ou responsáveis pela licitação, nos termos do Art° 9, Inciso III, da Lei n°"
        " 8.666/93.")
    texto_declaracao_05.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DECLARAÇÃO 06
    texto_declaracao_06 = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao_06.add_run("DECLARAMOS").bold = True
    texto_declaracao_06.add_run(
        " que nossa empresa se compromete a realizar os serviços objeto do presente edital em conformidade"
        " com a legislação ambiental vigente.")
    texto_declaracao_06.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA.", style="TextoSemIdentacao")
    linha_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = documento.add_paragraph("{}".format(nomeRepresentanteLegal), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = documento.add_paragraph("RG: {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor),
                                       style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    documento.save("{}/(06) DECLARAÇÕES DIVERSAS.docx".format(enderecoSalvar))


def criar_documento07_execucao_de_acordo(razaoSocial, CNPJ, concorrencia, objeto, data, nomeRepresentanteLegal,
                                         identidadeRepresentanteLegal, orgaoEmissor, enderecoSalvar):
    documento = DocumentMultivendas.criarComEstilo()

    # TITULO DO DOCUMENTO
    titulo = documento.add_paragraph("DECLARAÇÃO DE EXECUÇÃO CONFORME O PROJETO", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO CABEÇALHO DOCUMENTO 07
    texto_declaracao_cabecalho = documento.add_paragraph("A empresa", style="TextoAnexos")
    texto_declaracao_cabecalho.add_run(" {}".format(razaoSocial)).bold = True
    texto_declaracao_cabecalho.add_run(", CNPJ {}".format(CNPJ))
    texto_declaracao_cabecalho.add_run(", em atendimento ao Edital da")
    texto_declaracao_cabecalho.add_run(" {}".format(concorrencia)).bold = True
    texto_declaracao_cabecalho.add_run(", tendo como objeto:")
    texto_declaracao_cabecalho.add_run(" {}".format(objeto)).bold = True
    texto_declaracao_cabecalho.add_run(", vem declarar conforme abaixo:")
    texto_declaracao_cabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DECLARAÇÃO PROPRIAMENTE DITA
    texto_declaracao = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao.add_run("DECLARAMOS").bold = True
    texto_declaracao.add_run(" que executaremos o(s) serviço(s) de acordo com os projetos, especificações técnicas e"
                             " planilha orçamentária, assim como alocaremos os equipamentos, pessoal técnico especializado e"
                             " materiais necessários. Também declaramos que tomaremos todas as medidas para assegurar um controle adequado da "
                             "qualidade, prevenção e mitigação dos impactos sobre o meio ambiente, sobre os usuários e moradores"
                             " vizinhos.")
    texto_declaracao.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA.", style="TextoSemIdentacao")
    linha_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = documento.add_paragraph("{}".format(nomeRepresentanteLegal), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = documento.add_paragraph("RG: {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor),
                                       style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    documento.save("{}/(07) DECLARAÇÃO DE EXECUÇÃO CONFORME O PROJETO.docx".format(enderecoSalvar))


def criar_documento08_pleno_conhecimento(razaoSocial, CNPJ, concorrencia, objeto, data, nomeRepresentanteLegal,
                                         identidadeRepresentanteLegal, orgaoEmissor, enderecoSalvar):
    documento = DocumentMultivendas.criarComEstilo()

    # TITULO DO DOCUMENTO
    titulo = documento.add_paragraph("DECLARAÇÃO DE PLENO CONHECIMENTO", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO CABEÇALHO DOCUMENTO 08
    texto_declaracao_cabecalho = documento.add_paragraph("A empresa", style="TextoAnexos")
    texto_declaracao_cabecalho.add_run(" {}".format(razaoSocial)).bold = True
    texto_declaracao_cabecalho.add_run(", CNPJ {}".format(CNPJ))
    texto_declaracao_cabecalho.add_run(", em atendimento ao Edital da")
    texto_declaracao_cabecalho.add_run(" {}".format(concorrencia)).bold = True
    texto_declaracao_cabecalho.add_run(", tendo como objeto:")
    texto_declaracao_cabecalho.add_run(" {}".format(objeto)).bold = True
    texto_declaracao_cabecalho.add_run(", vem declarar conforme abaixo:")
    texto_declaracao_cabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DECLARAÇÃO PROPRIAMENTE DITA
    texto_declaracao = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao.add_run("DECLARAMOS").bold = True
    texto_declaracao.add_run(
        " que temos pleno conhecimento das condições e peculiaridades existentes inerentes a natureza"
        " dos serviços a serem executados (condições dos locais para a execução do objeto), assim como"
        " assumimos total responsabilidade por esta declaração, consequentemente, ficamos impedidos de, no futuro,"
        " pleitear qualquer desconhecimento do local e alterações contratuais de natureza técnica"
        " e/ou financeira.")
    texto_declaracao.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA.", style="TextoSemIdentacao")
    linha_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = documento.add_paragraph("{}".format(nomeRepresentanteLegal), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = documento.add_paragraph("RG: {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor),
                                       style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    documento.save("{}/(08) DECLARAÇÃO DE PLENO CONHECIMENTO.docx".format(enderecoSalvar))


def criar_documento09_recebimento_edital(razaoSocial, CNPJ, concorrencia, objeto, data, nomeRepresentanteLegal,
                                         identidadeRepresentanteLegal, orgaoEmissor, enderecoSalvar):
    documento = DocumentMultivendas.criarComEstilo()

    # TITULO DO DOCUMENTO
    titulo = documento.add_paragraph("DECLARAÇÃO DE RECEBIMENTO DE TODAS AS PEÇAS RELATIVAS A LICITAÇÃO",
                                     style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO CABEÇALHO DOCUMENTO 09
    texto_declaracao_cabecalho = documento.add_paragraph("A empresa", style="TextoAnexos")
    texto_declaracao_cabecalho.add_run(" {}".format(razaoSocial)).bold = True
    texto_declaracao_cabecalho.add_run(", CNPJ {}".format(CNPJ))
    texto_declaracao_cabecalho.add_run(", em atendimento ao Edital da")
    texto_declaracao_cabecalho.add_run(" {}".format(concorrencia)).bold = True
    texto_declaracao_cabecalho.add_run(", tendo como objeto:")
    texto_declaracao_cabecalho.add_run(" {}".format(objeto)).bold = True
    texto_declaracao_cabecalho.add_run(", vem declarar conforme abaixo:")
    texto_declaracao_cabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DECLARAÇÃO PROPRIAMENTE DITA
    texto_declaracao = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao.add_run("DECLARAMOS").bold = True
    texto_declaracao.add_run(" que recebemos todas as peças relativas à licitação, as quais são: Edital, orçamentos,"
                             " cronogramas, memoriais de cálculo, especificações, plantas gráficas e outras materiais"
                             " pertinentes à licitação.")
    texto_declaracao.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA.", style="TextoSemIdentacao")
    linha_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = documento.add_paragraph("{}".format(nomeRepresentanteLegal), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = documento.add_paragraph("RG: {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor),
                                       style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    documento.save("{}/(09) DECLARAÇÃO DE RECEBIMENTO DO EDITAL.docx".format(enderecoSalvar))


def criar_documento10_ensaios_tecnologicos(razaoSocial, CNPJ, concorrencia, objeto, data, nomeRepresentanteLegal,
                                           identidadeRepresentanteLegal, orgaoEmissor, enderecoSalvar):
    documento = DocumentMultivendas.criarComEstilo()

    # TEXTO CABEÇALHO DOCUMENTO 10
    texto_declaracao_cabecalho = documento.add_paragraph("A empresa", style="TextoAnexos")
    texto_declaracao_cabecalho.add_run(" {}".format(razaoSocial)).bold = True
    texto_declaracao_cabecalho.add_run(", CNPJ {}".format(CNPJ))
    texto_declaracao_cabecalho.add_run(", em atendimento ao Edital da")
    texto_declaracao_cabecalho.add_run(" {}".format(concorrencia)).bold = True
    texto_declaracao_cabecalho.add_run(", tendo como objeto:")
    texto_declaracao_cabecalho.add_run(" {}".format(objeto)).bold = True
    texto_declaracao_cabecalho.add_run(", vem declarar conforme abaixo:")
    texto_declaracao_cabecalho.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DECLARAÇÃO PROPRIAMENTE DITA
    texto_declaracao = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao.add_run("DECLARAMOS").bold = True
    texto_declaracao.add_run(" que apresentaremos à fiscalização relatórios consubstanciados com dados essenciais"
                             " dos levantamentos e ensaios tecnológicos para a avaliação da qualidade dos serviços"
                             " executados em suas diversas fases, sempre que se fizer necessário ou de acordo com"
                             " previsão no projeto/medição dos serviços.")
    texto_declaracao.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    texto_declaracao_02 = documento.add_paragraph("", style="TextoAnexos")
    texto_declaracao_02.add_run("DECLARAMOS").bold = True
    texto_declaracao_02.add_run(" que executaremos o controle tecnológico conforme previsto nas especificações"
                                " técnicas e normas técnicas relacionadas sempre que for solicitado pela fiscalização.")
    texto_declaracao_02.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA.", style="TextoSemIdentacao")
    linha_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = documento.add_paragraph("{}".format(nomeRepresentanteLegal), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = documento.add_paragraph("RG: {} ({})".format(identidadeRepresentanteLegal, orgaoEmissor),
                                       style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    documento.save("{}/(10) DECLARAÇÃO DE ENSAIOS TECNOLÓGICOS.docx".format(enderecoSalvar))


def criar_documento11_informacoes_contrato(data, enderecoSalvar):
    documento = DocumentMultivendas.criarComEstiloSemRodape()

    # TITULO DO DOCUMENTO 11 (INFORMAÇÕES PARA FORMALIZAÇÃO DO CONTRATO)
    titulo = documento.add_paragraph("INFORMAÇÕES PARA FORMALIZAÇÃO DO CONTRATO", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # DADOS
    titulo_dados_empresa = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_dados_empresa.add_run("1 - DADOS DA EMPRESA:").bold = True

    titulo_razao_social = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_razao_social.add_run("RAZAO SOCIAL:").bold = True
    titulo_razao_social.add_run(" MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA.")

    titulo_nome_fantasia = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_nome_fantasia.add_run("NOME FANTAZIA: ").bold = True
    titulo_nome_fantasia.add_run(" MULTIVENDAS")

    titulo_cnpj = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_cnpj.add_run("CNPJ:").bold = True
    titulo_cnpj.add_run(" 07.538.900/0001-36")

    titulo_endereco = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_endereco.add_run("ENDEREÇO:").bold = True
    titulo_endereco.add_run(" Rua Cecília Brasil, 1274")

    titulo_endereco_multiplo = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_endereco_multiplo.add_run("BAIRRO: ").bold = True
    titulo_endereco_multiplo.add_run(" CENTRO")
    titulo_endereco_multiplo.add_run(" CIDADE:").bold = True
    titulo_endereco_multiplo.add_run(" BOA VISTA")
    titulo_endereco_multiplo.add_run(" ESTADO:").bold = True
    titulo_endereco_multiplo.add_run(" RORAIMA")

    titulo_cep = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_cep.add_run("CEP: ").bold = True
    titulo_cep.add_run(" 69.301-080")

    titulo_fone_contato = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_fone_contato.add_run("FONE CONTATO: ").bold = True
    titulo_fone_contato.add_run(" (95) 3623-0720")
    titulo_fone_contato.add_run("   CEL:").bold = True
    titulo_fone_contato.add_run(" (95) 99112-6407")

    titulo_email = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_email.add_run("E-MAIL:").bold = True
    titulo_email.add_run(" multivendas2@yahoo.com.br")

    titulo_banco = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_banco.add_run("BANCO:").bold = True
    titulo_banco.add_run(" 341-ITAÚ -")
    titulo_banco.add_run(" AGÊNCIA:").bold = True
    titulo_banco.add_run(" 8526")
    titulo_banco.add_run(" C/C:").bold = True
    titulo_banco.add_run(" 8551-2")

    titulo_nome_contato = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_nome_contato.add_run("NOME PARA CONTATO:").bold = True
    titulo_nome_contato.add_run("JOSEILDO SOARES DE SOUZA")

    titulo_dados_representante = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_dados_representante.add_run("\n\n2 - ASSINATURA DO CONTRATO").bold = True

    titulo_nome_representante_legal = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_nome_representante_legal.add_run("NOME:").bold = True
    titulo_nome_representante_legal.add_run(" JOSEILDO SOARES DE SOUZA")

    titulo_nacionalidade = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_nacionalidade.add_run("NACIONALIDADE:").bold = True
    titulo_nacionalidade.add_run(" BRASILEIRO")

    titulo_naturalidade = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_naturalidade.add_run("NATURALIDADE:").bold = True
    titulo_naturalidade.add_run(" PERNAMBUCANO")

    titulo_profissao = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_profissao.add_run("PROFISSÃO:").bold = True
    titulo_profissao.add_run(" EMPRESÁRIO")

    titulo_estado_civil = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_estado_civil.add_run("ESTADO CIVIL:").bold = True
    titulo_estado_civil.add_run(" SOLTEIRO")

    titulo_endereco_resid = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_endereco_resid.add_run("ENDEREÇO RESID:").bold = True
    titulo_endereco_resid.add_run(" AV. BRIG. EDUARDO GOMES, 346 - DOS ESTADOS")

    titulo_ci = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_ci.add_run("CI:").bold = True
    titulo_ci.add_run(" 1973885/SSP – PA")

    titulo_cpf = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_cpf.add_run("CPF:").bold = True
    titulo_cpf.add_run(" 329.171.202-15")

    titulo_na_qualidade = documento.add_paragraph("", style="SubTitleAnexos")
    titulo_na_qualidade.add_run("NA QUALIDADE DE:").bold = True
    titulo_na_qualidade.add_run(" REPRESENTANTE LEGAL (SÓCIO)")

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("\n\n\n\nBoa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA.", style="TextoSemIdentacao")
    linha_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = documento.add_paragraph("JOSEILDO SOARES DE SOUZA", style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = documento.add_paragraph("RG: 1973885 SSP – PA",
                                       style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    documento.save("{}/(11) INFORMAÇÕES PARA CONTRATO.docx".format(enderecoSalvar))


def criar_documento_carta_proposta(concorrencia, processo, data_abertura, horario_abertura, objeto,
                                   valor_total_proposta,
                                   prazo_validade_proposta, prazo_garantia, prazo_de_execucao_contrato,
                                   nome_responsavel_tecnico, identidade_responsavel_tecnico, orgao_emissor,
                                   endereco_salvar):
    documento = DocumentMultivendas.criarComEstilo()

    documento.add_paragraph("AO MUNICÍPIO DE BOA VISTA", style="SubTitleAnexos")
    paragrafo_conc_proc = documento.add_paragraph("Ref.: {} - {}".format(concorrencia, processo), style="SubTitleAnexos")
    paragrafo_conc_proc.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    documento.add_paragraph("", style="SubTitleAnexos").add_run("DIA DE ABERTURA: {}".format(data_abertura)).bold = True
    documento.add_paragraph("", style="SubTitleAnexos").add_run("HORÁRIO: {}".format(horario_abertura)).bold = True
    documento.add_paragraph("\nPrezados senhores,", style="SubTitleAnexos")
    documento.add_paragraph("\nEncaminhamos a V.S.ª nossas propostas para: ", style="SubTitleAnexos")

    paragrafo_objeto = documento.add_paragraph("{}".format(objeto), style="SubTitleAnexos")
    paragrafo_objeto.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragrafo_valor_total = documento.add_paragraph("", style="SubTitleAnexos")
    paragrafo_valor_total.add_run("VALOR TOTAL: R$ {}".format(valor_total_proposta)).bold = True
    numero_por_extenso = number_to_long_number(valor_total_proposta)
    numero_por_extenso = capitalizar_letras(numero_por_extenso)
    paragrafo_valor_total.add_run(" ({})".format(numero_por_extenso))
    paragrafo_valor_total.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    documento.add_paragraph("Outrossim, declaramos que:", style="Alinea").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    documento.add_paragraph("a) Temos conhecimento do local onde será executada a obra;", style="Alinea").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    documento.add_paragraph("b) Aceitamos todas as condições impostas pelo Edital e seus anexos;", style="Alinea").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    documento.add_paragraph("c) As obras serão executadas e concluídas dentro do prazo fixado no Edital;", style = "Alinea").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    documento.add_paragraph("d) Esta proposta compreende todas as despesas com mão de obra (inclusive leis sociais)"
                            ", materiais, ferramentas, transportes, equipamentos, seguros, impostos, BDI e"
                            " demais encargos necessários à perfeita execução de toda a obra;", style = "Alinea").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragrafo_validade_proposta = documento.add_paragraph("e) Concordamos em manter a validade desta proposta por "
                                                          "um período de", style = "Alinea")
    paragrafo_validade_proposta.add_run(" {}".format(prazo_validade_proposta)).bold = True
    paragrafo_validade_proposta.add_run(", contado da data final prevista para sua entrega.")
    paragrafo_validade_proposta.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragrafo_garantia = documento.add_paragraph("f) O prazo de garantia é de", style="Alinea")
    paragrafo_garantia.add_run(" {}".format(prazo_garantia)).bold = True
    paragrafo_garantia.add_run(", a contar da data de recebimento definitivo de todos os serviços prestados"
                               " contra quaisquer defeitos de funcionamento das instalações e materiais"
                               " fornecidos;")
    paragrafo_garantia.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    documento.add_paragraph("g) Declaramos que todos os materiais a serem utilizados são de primeira qualidade"
                            " - 'Classe A' e estão de conformidade com as normas estabelecidas pela ABNT,"
                            " com certificação pelo INMETRO ou Normas ISO;", style="Alinea").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    paragrafo_execucao_contrato = documento.add_paragraph("h) Prazo de execução do contrato: ", style="Alinea")
    paragrafo_execucao_contrato.add_run(" {}".format(prazo_de_execucao_contrato)).bold = True
    paragrafo_execucao_contrato.add_run(", contados da emissão da Ordem de Serviço emitida pela CONTRATANTE, podendo ser"
                                        " prorrogado, nos termos do Art. 57, §1° e incisos, e §2° da Lei 8.666/93.")
    paragrafo_execucao_contrato.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    documento.add_paragraph("Até que o contrato seja assinado, esta proposta constituirá um compromisso de nossa parte,"
                            " observadas as condições de Edital.", style = "TextoAnexos").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    documento.add_paragraph("Responsável legal e responsável técnico: Eng. Civil. Joseildo Soares de Souza;", style="Alinea").alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    documento.add_paragraph("Banco ITAÚ, Agência 8526, Conta Corrente n° 8551-2", style="Alinea")
    documento.add_paragraph("Atenciosamente, ", style="SubTitleAnexos")

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data_abertura), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA
    linha_assinatura = documento.add_paragraph("___________________________________", style="TextoSemIdentacao")
    linha_assinatura.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA.", style="TextoSemIdentacao")
    linha_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE
    linha_representante_legal = documento.add_paragraph("{}".format(nome_responsavel_tecnico), style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # RG DO REPRESENTANTE LEGAL
    linha_rg = documento.add_paragraph("RG: {} ({})".format(identidade_responsavel_tecnico, orgao_emissor),
                                       style="TextoSemIdentacao")
    linha_rg.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    linha_rg.paragraph_format.line_spacing = Inches(CM2IN(0.50))

    documento.save("{}/(02) CARTA PROPOSTA.docx".format(endereco_salvar))

def criar_declaracao_independente_proposta(razaoSocial, CNPJ, concorrencia, processo, data, nomeRepresentanteLegal,
                                           identidadeRepresentanteLegal, orgaoEmissor, cpfRepresentanteLegal,
                                           enderecoSalvar):

    documento = DocumentMultivendas.criarComEstilo()

    # TITULO DO DOCUMENTO (DECLARAÇÃO DE ELABORAÇÃO INDEPENDENTE DE PROPOSTA)
    titulo = documento.add_paragraph("DECLARAÇÃO DE ELABORAÇÃO INDEPENDENTE DE PROPOSTA", style="TitleAnexos")
    titulo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    subtitulo_concorrencia = documento.add_paragraph("{}".format(concorrencia), style="SubTitleAnexos")
    subtitulo_concorrencia.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    subtitulo_processo = documento.add_paragraph("{}".format(processo), style="SubTitleAnexos")
    subtitulo_processo.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # TEXTO PRINCIPAL DO DOCUMENTO 04
    texto_principal = documento.add_paragraph("A", style="TextoAnexos")
    texto_principal.add_run(" {}".format(razaoSocial)).bold = True
    texto_principal.add_run(", inscrita no CNPJ n° {}, doravante denominada (Licitante), por intermédio de seu"
                            " representante legal o(a) Sr.(a) {}, portador(a) da Carteira"
                            " de Identidade n°{}/{} e do CPF n° {}, para fins do disposto no Edital da"
                            .format(CNPJ, nomeRepresentanteLegal, identidadeRepresentanteLegal, orgaoEmissor,
                                    cpfRepresentanteLegal))
    texto_principal.add_run(" {}".format(concorrencia)).bold = True
    texto_principal.add_run(", DECLARA, sob as penas da lei, em especial o art. 299 do Código Penal Brasileiro, que: "
                            "a proposta apresentada foi elaborada de maneira independente (pelo Licitante), e que o "
                            "conteúdo da proposta anexa não foi, no todo ou em parte, direta ou indiretamente, "
                            "informado, discutido com ou recebido de qualquer outro participante potencial ou de fato do"
                            " certame em referência, por qualquer meio ou por qualquer pessoa; a intenção de apresentar "
                            "a proposta elaborada não foi informada, ou discutida com ou recebida de qualquer outro "
                            "participante potencial ou de fato do certame em referência, por qualquer meio ou por "
                            "qualquer pessoa; que não tentou, por qualquer meio ou por qualquer pessoa, influir na "
                            "decisão de qualquer outro participante potencial ou de fato do certame em referência, "
                            "quanto a participar ou não da referida licitação; que o conteúdo da proposta anexa não "
                            "será, no todo ou em parte, direta ou indiretamente, comunicado a ou discutido com qualquer "
                            "outro participante potencial ou de fato do certame em referência, antes da adjudicação do "
                            "objeto da referida licitação; que o conteúdo da proposta anexa não foi, no todo ou em "
                            "parte, direta ou indiretamente, informado a, discutido com ou recebido de qualquer "
                            "integrante da Comissão Permanente de Licitação – CPL antes da abertura oficial das "
                            "propostas; e que está plenamente ciente do teor e da extensão desta declaração e que "
                            "detém plenos poderes e informações para firmá-la.")

    texto_principal.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # TEXTO DA DATA
    texto_data = documento.add_paragraph("Boa Vista - RR, {}".format(data), style='DataAnexos')
    texto_data.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

    # LINHA PARA ASSINATURA DO REPRESENTANTE LEGAL
    linha_assinatura_02 = documento.add_paragraph("\n___________________________________", style="TextoSemIdentacao")
    linha_assinatura_02.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DA EMPRESA
    linha_nome_empresa = documento.add_paragraph("MULTIVENDAS COMÉRCIO E SERVIÇOS LTDA", style="TextoSemIdentacao")
    linha_nome_empresa.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_nome_empresa.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # NOME DO REPRESENTANTE LEGAL
    linha_representante_legal = documento.add_paragraph("{}".format(nomeRepresentanteLegal),
                                                        style="TextoSemIdentacao")
    linha_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # CARGO DO REPRESENTANTE LEGAL
    linha_cargo_representante_legal = documento.add_paragraph("RG: {} {}".format(identidadeRepresentanteLegal,
                                                                                 orgaoEmissor),
                                                              style="TextoSemIdentacao")

    linha_cargo_representante_legal.paragraph_format.line_spacing = Inches(CM2IN(0.50))
    linha_cargo_representante_legal.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    documento.save("{}/(03) DECLARAÇÃO DE ELABORAÇÃO INDEPENDENTE DE PROPOSTA.docx".format(enderecoSalvar))

def criar_envelope_propostas(concorrencia, processo, objeto, data_abertura, horario_abertura, endereco_salvar):
    documento = Document("C:/Users/willi/PycharmProjects/MultivendasLicitacao/configs/(14) ENVELOPE PREFEITURA.docx")

    documento.paragraphs[8].text = concorrencia
    documento.paragraphs[9].text = processo
    documento.paragraphs[14].text = "OBJETO: {}".format(objeto)
    documento.paragraphs[16].text = "DATA DE ABERTURA: {} ÀS {}".format(data_abertura, horario_abertura)

    documento.paragraphs[26].text = concorrencia
    documento.paragraphs[27].text = processo
    documento.paragraphs[32].text = "OBJETO: {}".format(objeto)
    documento.paragraphs[34].text = "DATA DE ABERTURA: {} ÀS {}".format(data_abertura, horario_abertura)

    documento.paragraphs[53].text = concorrencia
    documento.paragraphs[54].text = processo
    documento.paragraphs[59].text = "OBJETO: {}".format(objeto)
    documento.paragraphs[61].text = "DATA DE ABERTURA: {} ÀS {}".format(data_abertura, horario_abertura)

    documento.save("{}/(14) ENVELOPES PREFEITURA.docx".format(endereco_salvar))


def criar_capa_cd_proposta(concorrencia, processo, objeto, enderecoSalvar):
    documento = DocumentMultivendas.criarComEstiloSemRodape()

    for section in documento.sections:
        section.different_first_page_header_footer = False
        section.header.is_linked_to_previous = True

    # parágrafo concorrencia
    paragrafo_concorrencia = documento.add_paragraph("{}".format(concorrencia), style="CD")

    # parágrafo processo
    paragrafo_processo = documento.add_paragraph("{}".format(processo), style="CD")

    # paragrafo objeto
    paragrafo_objeto = documento.add_paragraph("", style="CD")
    paragrafo_objeto.add_run("OBJETO: ").bold = True
    paragrafo_objeto.add_run("{}".format(objeto))
    paragrafo_objeto.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY

    # paragrafos itens
    paragrafo_contem = documento.add_paragraph("", style="CD")
    paragrafo_contem.add_run("CONTÉM:").bold = True
    documento.add_paragraph("1 - Proposta", style="CD")
    documento.add_paragraph("2 - Cronograma", style="CD")
    documento.add_paragraph("3 - Composições de preços unitários", style="CD")
    documento.add_paragraph("4 - Composição do BDI", style="CD")
    documento.add_paragraph("5 - Composição LS", style="CD")
    documento.add_paragraph("6 - Lista de insumos", style="CD")

    documento.save("{}/(01) CAPA CD DA PROPOSTA.docx".format(enderecoSalvar))
